from __future__ import annotations

import base64
from dataclasses import dataclass
from functools import lru_cache
import json
import logging
import os
from pathlib import Path
from typing import Any

from app.config import settings


VISUAL_EVIDENCE_DOCUMENT_TYPES = {"Photos of Damage", "Spare Key"}
PDF_MIME_TYPES = {"application/pdf"}
DOCX_MIME_TYPES = {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
TEXT_MIME_TYPES = {"text/plain"}
IMAGE_MIME_TYPES = {"image/jpeg", "image/png", "image/webp"}
IMAGE_SUFFIXES = {".jpg", ".jpeg", ".png", ".webp"}
IMAGE_MEDIA_TYPES = {".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".png": "image/png", ".webp": "image/webp"}
OPENROUTER_CHAT_COMPLETIONS_URL = "https://openrouter.ai/api/v1/chat/completions"
logger = logging.getLogger(__name__)


class DocumentExtractionError(ValueError):
    pass


@dataclass(frozen=True)
class ExtractedDocument:
    strategy: str
    text: str
    confidence: float | None
    structure: dict[str, Any] | None = None
    structured_data: dict[str, Any] | None = None


def extract_document_content(path: Path, mime_type: str, document_type: str) -> ExtractedDocument:
    if document_type in VISUAL_EVIDENCE_DOCUMENT_TYPES:
        return ExtractedDocument(strategy="visual_evidence_preserved", text="", confidence=None)

    suffix = path.suffix.lower()
    normalized_mime_type = (mime_type or "").lower()
    if normalized_mime_type in PDF_MIME_TYPES or suffix == ".pdf":
        text = _extract_pdf_text(path)
        if text:
            return _with_ai_structured_data(ExtractedDocument(strategy="native_pdf_text", text=text, confidence=1.0), document_type)
        return _with_ai_structured_data(_extract_scanned_pdf_text(path), document_type)
    if normalized_mime_type in DOCX_MIME_TYPES or suffix == ".docx":
        return _with_ai_structured_data(ExtractedDocument(strategy="native_docx_text", text=_extract_docx_text(path), confidence=1.0), document_type)
    if normalized_mime_type in TEXT_MIME_TYPES or suffix == ".txt":
        return _with_ai_structured_data(ExtractedDocument(strategy="native_text", text=_read_plain_text(path), confidence=1.0), document_type)
    if normalized_mime_type in IMAGE_MIME_TYPES or suffix in IMAGE_SUFFIXES:
        return _with_ai_structured_data(_extract_image_text(path, strategy="image_ocr"), document_type)
    raise DocumentExtractionError("This file type is not supported for document content extraction")


def _extract_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(path)
        return _normalise_text("\n".join(page.extract_text() or "" for page in reader.pages))
    except Exception as exc:
        raise DocumentExtractionError("The PDF could not be read") from exc


def _extract_scanned_pdf_text(path: Path) -> ExtractedDocument:
    try:
        import pypdfium2 as pdfium

        document = pdfium.PdfDocument(path)
        images = []
        for page in document:
            bitmap = page.render(scale=2)
            images.append(bitmap.to_pil())
        document.close()
    except Exception as exc:
        raise DocumentExtractionError("The scanned PDF could not be rendered") from exc
    try:
        # Retain the established strategy string for API consumers.
        return _extract_images_with_ocr(images, strategy="scanned_pdf_ocr")
    finally:
        for image in images:
            image.close()


def _extract_image_text(path: Path, strategy: str) -> ExtractedDocument:
    try:
        logger.info("Trying OpenRouter OCR for image")
        extracted = _openrouter_image_ocr(path, strategy)
        logger.info("OpenRouter OCR succeeded")
        return extracted
    except Exception as exc:
        logger.warning("OpenRouter OCR failed, falling back to Groq: %s", exc)
    try:
        logger.info("Trying Groq OCR fallback")
        extracted = _groq_image_ocr(path, strategy)
        logger.info("Groq OCR succeeded")
        return extracted
    except Exception as fallback_exc:
        logger.exception("Both OCR providers failed")
        raise DocumentExtractionError("OCR could not read the image") from fallback_exc


def _openrouter_image_ocr(path: Path, strategy: str) -> ExtractedDocument:
    api_key = settings.OPENROUTER_API_KEY
    model = settings.OPENROUTER_VISION_MODEL
    if not api_key or not model:
        raise DocumentExtractionError("OpenRouter OCR is not configured")
    payload = _vision_request_payload(path, model)
    response = _openrouter_http_client().post(
        OPENROUTER_CHAT_COMPLETIONS_URL,
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        json=payload,
    )
    response.raise_for_status()
    try:
        content = response.json()["choices"][0]["message"]["content"]
    except (KeyError, IndexError, TypeError, ValueError) as exc:
        raise DocumentExtractionError("OpenRouter returned an invalid OCR response") from exc
    return _vision_response_to_extracted_document(content, strategy, "openrouter")


@lru_cache(maxsize=1)
def _openrouter_http_client():
    import httpx

    return httpx.Client(timeout=settings.OCR_API_TIMEOUT_SECONDS)


def _groq_image_ocr(path: Path, strategy: str) -> ExtractedDocument:
    api_key = settings.GROQ_API_KEY
    model = settings.GROQ_VISION_MODEL
    if not api_key or not model:
        raise DocumentExtractionError("Groq OCR is not configured")
    completion = _groq_client().chat.completions.create(**_vision_request_payload(path, model))
    try:
        content = completion.choices[0].message.content
    except (AttributeError, IndexError, TypeError) as exc:
        raise DocumentExtractionError("Groq returned an invalid OCR response") from exc
    return _vision_response_to_extracted_document(content, strategy, "groq")


@lru_cache(maxsize=1)
def _groq_client():
    if not settings.GROQ_API_KEY:
        raise DocumentExtractionError("Groq OCR is not configured")
    try:
        from groq import Groq

        return Groq(api_key=settings.GROQ_API_KEY, timeout=settings.OCR_API_TIMEOUT_SECONDS)
    except Exception as exc:
        raise DocumentExtractionError("Groq OCR client could not be initialized") from exc


def _with_ai_structured_data(extracted: ExtractedDocument, document_type: str) -> ExtractedDocument:
    """Attach semantic data without allowing provider failures to discard extraction."""
    if extracted.structured_data:
        return extracted
    try:
        structured_data = _extract_structured_data_with_ai(extracted.text, document_type)
    except Exception:
        logger.exception("AI structured extraction failed; retaining the existing extraction result")
        return extracted
    return ExtractedDocument(
        strategy=extracted.strategy,
        text=extracted.text,
        confidence=extracted.confidence,
        structure=extracted.structure,
        structured_data=structured_data,
    )


def _extract_structured_data_with_ai(text: str, document_type: str) -> dict[str, Any] | None:
    if not text.strip():
        return None
    try:
        logger.info("Trying OpenRouter structured extraction")
        result = _openrouter_structured_extraction(text, document_type)
        logger.info("OpenRouter structured extraction succeeded")
        return result
    except Exception as exc:
        logger.warning("OpenRouter structured extraction failed, falling back to Groq: %s", exc)
    try:
        logger.info("Trying Groq structured extraction fallback")
        result = _groq_structured_extraction(text, document_type)
        logger.info("Groq structured extraction succeeded")
        return result
    except Exception as exc:
        logger.warning("Both structured extraction providers failed: %s", exc)
        return None


def _openrouter_structured_extraction(text: str, document_type: str) -> dict[str, Any] | None:
    if not settings.OPENROUTER_API_KEY or not settings.OPENROUTER_VISION_MODEL:
        raise DocumentExtractionError("OpenRouter structured extraction is not configured")
    response = _openrouter_http_client().post(
        OPENROUTER_CHAT_COMPLETIONS_URL,
        headers={"Authorization": f"Bearer {settings.OPENROUTER_API_KEY}", "Content-Type": "application/json"},
        json=_structured_request_payload(text, document_type, settings.OPENROUTER_VISION_MODEL),
    )
    response.raise_for_status()
    try:
        content = response.json()["choices"][0]["message"]["content"]
    except (KeyError, IndexError, TypeError, ValueError) as exc:
        raise DocumentExtractionError("OpenRouter returned an invalid structured extraction response") from exc
    return _structured_data_from_response(content)


def _groq_structured_extraction(text: str, document_type: str) -> dict[str, Any] | None:
    if not settings.GROQ_API_KEY or not settings.GROQ_VISION_MODEL:
        raise DocumentExtractionError("Groq structured extraction is not configured")
    completion = _groq_client().chat.completions.create(**_structured_request_payload(text, document_type, settings.GROQ_VISION_MODEL))
    try:
        content = completion.choices[0].message.content
    except (AttributeError, IndexError, TypeError) as exc:
        raise DocumentExtractionError("Groq returned an invalid structured extraction response") from exc
    return _structured_data_from_response(content)


def _structured_request_payload(text: str, document_type: str, model: str) -> dict[str, Any]:
    return {
        "model": model,
        "temperature": 0,
        "response_format": {"type": "json_object"},
        "messages": [{"role": "user", "content": _STRUCTURED_EXTRACTION_PROMPT.format(document_type=document_type, text=text)}],
    }


_STRUCTURED_EXTRACTION_PROMPT = """Act as a document-understanding and semantic information-extraction system. Return only valid JSON, with no markdown or explanation. Analyze the actual content; do not blindly trust the supplied document type when it conflicts with the content. Preserve names, numbers, IDs, dates, amounts, and values exactly; do not translate, summarize, infer, or hallucinate. Use null for unavailable values. Return semantic fields relevant to the document rather than copying the raw text. Use exactly this wrapper: {{\"detected_document_type\": \"...\", \"structured_data\": {{...}}}}. Incoming document type: {document_type}\n\nExtracted document text:\n{text}"""


def _structured_data_from_response(content: Any) -> dict[str, Any] | None:
    payload = _parse_model_json(content, "structured extraction")
    structured_data = payload.get("structured_data")
    return structured_data if isinstance(structured_data, dict) and structured_data else None


def _vision_request_payload(path: Path, model: str) -> dict[str, Any]:
    media_type = IMAGE_MEDIA_TYPES.get(path.suffix.lower())
    if media_type is None:
        raise DocumentExtractionError("This image type is not supported for OCR")
    try:
        image_data = base64.b64encode(path.read_bytes()).decode("ascii")
    except OSError as exc:
        raise DocumentExtractionError("The image could not be read") from exc
    return {
        "model": model,
        "temperature": 0,
        "response_format": {"type": "json_object"},
        "messages": [{"role": "user", "content": [
            {"type": "text", "text": _VISION_OCR_PROMPT},
            {"type": "image_url", "image_url": {"url": f"data:{media_type};base64,{image_data}"}},
        ]}],
    }


_VISION_OCR_PROMPT = """Extract this document exactly. Read all visible Arabic, English, and mixed Arabic-English text; preserve numbers, IDs, dates, amounts, policy and claim numbers, forms, invoices, insurance documents, tables, sections, and logical reading order. Do not summarize, translate, explain, or invent content. Return only a JSON object with this shape: {\"text\": \"complete extracted text\", \"document_type\": \"detected document type\", \"blocks\": [{\"type\": \"text\", \"text\": \"...\", \"confidence\": null, \"bbox\": null, \"language\": \"arabic|english|mixed|unknown\"}], \"layout\": [{\"type\": \"paragraph|title|table|form|other\", \"text\": \"...\"}], \"tables\": [{\"html\": \"...\", \"rows\": []}], \"structured_data\": {}}. Return semantic structured_data based only on visible content. Use null for unavailable confidence or bounding boxes; never fabricate coordinates."""


def _vision_response_to_extracted_document(content: Any, strategy: str, provider: str) -> ExtractedDocument:
    if isinstance(content, list):
        content = "".join(item.get("text", "") if isinstance(item, dict) else str(item) for item in content)
    if not isinstance(content, str):
        raise DocumentExtractionError(f"{provider} returned an invalid OCR response")
    payload = _parse_model_json(content, "OCR")
    blocks = _vision_blocks(payload.get("blocks"), provider)
    text = _normalise_text(str(payload.get("text") or ""))
    if not text:
        text = _normalise_text("\n".join(block["text"] for block in blocks))
    if not text:
        raise DocumentExtractionError(f"{provider} returned empty OCR text")
    if not blocks:
        blocks = [{"type": "text", "text": text, "confidence": None, "bbox": None, "language_pass": provider}]
    confidences = [block["confidence"] for block in blocks if block["confidence"] is not None]
    structured_data = payload.get("structured_data")
    return ExtractedDocument(strategy=strategy, text=text, confidence=sum(confidences) / len(confidences) if confidences else None, structure={
        "version": 1,
        "reading_order": "top_to_bottom_left_to_right",
        "pages": [{"page": 1, "orientation": None, "blocks": blocks, "layout": _vision_layout(payload.get("layout")), "tables": _vision_tables(payload.get("tables"))}],
    }, structured_data=structured_data if isinstance(structured_data, dict) and structured_data else None)


def _parse_model_json(content: Any, response_name: str) -> dict[str, Any]:
    if isinstance(content, list):
        content = "".join(item.get("text", "") if isinstance(item, dict) else str(item) for item in content)
    if not isinstance(content, str):
        raise DocumentExtractionError(f"{response_name} provider returned an invalid response")
    candidate = content.strip()
    if candidate.startswith("```"):
        candidate = candidate.split("\n", 1)[-1].rsplit("```", 1)[0].strip()
    try:
        parsed = json.loads(candidate)
    except json.JSONDecodeError as exc:
        raise DocumentExtractionError(f"{response_name} provider returned malformed JSON") from exc
    if not isinstance(parsed, dict):
        raise DocumentExtractionError(f"{response_name} provider returned malformed JSON")
    return parsed


def _vision_blocks(value: Any, provider: str) -> list[dict[str, Any]]:
    if not isinstance(value, list):
        return []
    blocks = []
    for item in value:
        if not isinstance(item, dict) or not str(item.get("text") or "").strip():
            continue
        blocks.append({"type": str(item.get("type") or "text"), "text": str(item["text"]).strip(), "confidence": _float_or_none(item.get("confidence")), "bbox": _number_list(item.get("bbox")) or None, "language_pass": str(item.get("language") or provider)})
    return blocks


def _vision_layout(value: Any) -> list[dict[str, Any]]:
    return [item for item in value if isinstance(item, dict)] if isinstance(value, list) else []


def _vision_tables(value: Any) -> list[dict[str, Any]]:
    return [item for item in value if isinstance(item, dict)] if isinstance(value, list) else []


def _extract_docx_text(path: Path) -> str:
    try:
        from docx import Document

        document = Document(path)
        rows = [paragraph.text for paragraph in document.paragraphs]
        rows.extend(" | ".join(cell.text for cell in row.cells) for table in document.tables for row in table.rows)
        text = _normalise_text("\n".join(rows))
    except Exception as exc:
        raise DocumentExtractionError("The DOCX file could not be read") from exc
    if not text:
        raise DocumentExtractionError("The DOCX file contains no extractable text")
    return text


def _read_plain_text(path: Path) -> str:
    try:
        text = _normalise_text(path.read_text(encoding="utf-8"))
    except UnicodeDecodeError as exc:
        raise DocumentExtractionError("The text file is not UTF-8 encoded") from exc
    except OSError as exc:
        raise DocumentExtractionError("The text file could not be read") from exc
    if not text:
        raise DocumentExtractionError("The text file contains no extractable text")
    return text


@lru_cache(maxsize=1)
def _ocr_pipelines():
    """Load PP-StructureV3 once per language, with layout processing enabled.

    PP-OCRv5 is used for English. PaddleOCR's current Arabic recognition model
    is PP-OCRv3, so its supplementary pass preserves Arabic and mixed-language
    documents instead of relying on the English model for Arabic text.
    """
    _disable_onednn()
    from paddleocr import PPStructureV3

    shared_options = {
        "enable_mkldnn": False,
        "use_doc_orientation_classify": True,
        "use_doc_unwarping": True,
        "use_textline_orientation": True,
        "use_table_recognition": True,
        "use_formula_recognition": False,
        "use_seal_recognition": False,
    }
    return (
        ("english", PPStructureV3(lang="en", ocr_version="PP-OCRv5", **shared_options)),
        ("arabic", PPStructureV3(lang="ar", ocr_version="PP-OCRv3", **shared_options)),
    )


@lru_cache(maxsize=1)
def _text_ocr_pipelines():
    """Load the lighter text-only OCR pipelines used after a layout failure."""
    _disable_onednn()
    from paddleocr import PaddleOCR

    shared_options = {
        "enable_mkldnn": False,
        "use_doc_orientation_classify": True,
        "use_doc_unwarping": True,
        "use_textline_orientation": True,
    }
    return (
        ("english", PaddleOCR(lang="en", ocr_version="PP-OCRv5", **shared_options)),
        ("arabic", PaddleOCR(lang="ar", ocr_version="PP-OCRv3", **shared_options)),
    )


def _disable_onednn() -> None:
    """Set before importing PaddleOCR so each predictor is created without oneDNN."""
    os.environ["FLAGS_use_mkldnn"] = "0"


def _extract_images_with_ocr(images: list, strategy: str) -> ExtractedDocument:
    pages: list[dict[str, Any]] = []
    try:
        pipelines = _ocr_pipelines()
        for page_index, image in enumerate(images):
            try:
                page = _extract_page_structure(image, pipelines, page_index)
            except Exception:
                logger.exception("PPStructureV3 failed for image page %d; retrying with PaddleOCR text-only inference", page_index + 1)
                page = _extract_page_text_only(image, _text_ocr_pipelines(), page_index)
            pages.append(page)
    except Exception as exc:
        logger.exception("PaddleOCR failed while processing %d image page(s)", len(images))
        raise DocumentExtractionError("OCR could not read the document") from exc
    blocks = [block for page in pages for block in page["blocks"]]
    text = _normalise_text("\n".join(block["text"] for block in blocks))
    if not text:
        raise DocumentExtractionError("OCR could not find readable text in the document")
    confidences = [block["confidence"] for block in blocks if block["confidence"] is not None]
    structure = {
        "version": 1,
        "reading_order": "top_to_bottom_left_to_right",
        "pages": pages,
    }
    return ExtractedDocument(strategy=strategy, text=text, confidence=sum(confidences) / len(confidences) if confidences else None, structure=structure)


def _extract_page_structure(image, pipelines, page_index: int) -> dict[str, Any]:
    import numpy as np

    page_blocks: list[dict[str, Any]] = []
    layout: list[dict[str, Any]] = []
    tables: list[dict[str, Any]] = []
    orientation: int | None = None
    image_array = np.asarray(image.convert("RGB"))
    for language, pipeline in pipelines:
        for prediction in pipeline.predict(image_array):
            payload = _prediction_payload(prediction)
            if orientation is None:
                orientation = _orientation(payload)
            layout.extend(_layout_blocks(payload))
            tables.extend(_table_blocks(payload))
            page_blocks.extend(_ocr_blocks(payload, language))
    blocks = _deduplicate_and_sort_blocks(page_blocks)
    return {"page": page_index + 1, "orientation": orientation, "blocks": blocks, "layout": layout, "tables": tables}


def _extract_page_text_only(image, pipelines, page_index: int) -> dict[str, Any]:
    import numpy as np

    page_blocks: list[dict[str, Any]] = []
    orientation: int | None = None
    image_array = np.asarray(image.convert("RGB"))
    for language, pipeline in pipelines:
        for prediction in pipeline.predict(image_array):
            payload = _prediction_payload(prediction)
            if orientation is None:
                orientation = _orientation(payload)
            page_blocks.extend(_ocr_blocks(payload, language))
    return {
        "page": page_index + 1,
        "orientation": orientation,
        "blocks": _deduplicate_and_sort_blocks(page_blocks),
        "layout": [],
        "tables": [],
    }


def _prediction_payload(prediction: Any) -> dict[str, Any]:
    if isinstance(prediction, dict):
        payload = prediction
    else:
        value = getattr(prediction, "json", None)
        value = value() if callable(value) else value
        if isinstance(value, str):
            payload = json.loads(value)
        elif isinstance(value, dict):
            payload = value
        else:
            payload = dict(prediction)
    return payload.get("res", payload)


def _orientation(payload: dict[str, Any]) -> int | None:
    value = payload.get("doc_preprocessor_res", {}).get("angle")
    return int(value) if isinstance(value, (int, float)) else None


def _ocr_blocks(payload: dict[str, Any], language: str) -> list[dict[str, Any]]:
    result = payload.get("overall_ocr_res", {})
    texts = result.get("rec_texts", []) or []
    scores = result.get("rec_scores", []) or []
    polygons = result.get("rec_polys", result.get("dt_polys", [])) or []
    blocks = []
    for index, text in enumerate(texts):
        cleaned = str(text).strip()
        if not cleaned:
            continue
        score = scores[index] if index < len(scores) else None
        blocks.append({
            "type": "text",
            "text": cleaned,
            "confidence": float(score) if isinstance(score, (int, float)) else None,
            "bbox": _bbox(polygons[index]) if index < len(polygons) else None,
            "language_pass": language,
        })
    return blocks


def _layout_blocks(payload: dict[str, Any]) -> list[dict[str, Any]]:
    boxes = payload.get("layout_det_res", {}).get("boxes", []) or []
    return [
        {"type": box.get("label", "unknown"), "bbox": _number_list(box.get("coordinate")), "confidence": _float_or_none(box.get("score"))}
        for box in boxes if isinstance(box, dict)
    ]


def _table_blocks(payload: dict[str, Any]) -> list[dict[str, Any]]:
    tables = payload.get("table_res_list", payload.get("table_res", [])) or []
    if isinstance(tables, dict):
        tables = [tables]
    return [{"html": item.get("pred_html"), "bbox": _number_list(item.get("bbox"))} for item in tables if isinstance(item, dict)]


def _deduplicate_and_sort_blocks(blocks: list[dict[str, Any]]) -> list[dict[str, Any]]:
    unique: dict[tuple[str, tuple | None], dict[str, Any]] = {}
    for block in blocks:
        bbox = tuple(block["bbox"]) if block["bbox"] else None
        key = (block["text"], bbox)
        existing = unique.get(key)
        if existing is None or (block["confidence"] or 0) > (existing["confidence"] or 0):
            unique[key] = block
    return sorted(unique.values(), key=lambda block: ((block["bbox"] or [0, 0])[1], (block["bbox"] or [0, 0])[0]))


def _bbox(points: Any) -> list[float] | None:
    values = _number_list(points)
    if len(values) < 4:
        return None
    xs, ys = values[::2], values[1::2]
    return [min(xs), min(ys), max(xs), max(ys)]


def _number_list(value: Any) -> list[float]:
    if hasattr(value, "tolist"):
        value = value.tolist()
    if not isinstance(value, (list, tuple)):
        return []
    flattened = [item for group in value for item in group] if value and isinstance(value[0], (list, tuple)) else value
    return [float(item) for item in flattened if isinstance(item, (int, float))]


def _float_or_none(value: Any) -> float | None:
    return float(value) if isinstance(value, (int, float)) else None


def _normalise_text(text: str) -> str:
    return "\n".join(line.strip() for line in text.splitlines() if line.strip())
