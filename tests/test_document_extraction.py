import json
import os
from datetime import datetime
from pathlib import Path
from tempfile import TemporaryDirectory
from types import SimpleNamespace
from unittest.mock import patch
from uuid import uuid4
import unittest

from fastapi import HTTPException
from docx import Document
from PIL import Image

from app.document_extraction import DocumentExtractionError, extract_document_content
from app.main import extract_claim_document
from app.structured_extraction import extract_structured_data


class Result:
    def __init__(self, value):
        self.value = value

    def mappings(self):
        return self

    def first(self):
        return self.value

    def one(self):
        return self.value


class StubStructurePipeline:
    def __init__(self, payload):
        self.payload = payload
        self.inputs = []

    def predict(self, image):
        self.inputs.append(image)
        return [self.payload]


class FailingStructurePipeline:
    def predict(self, image):
        raise RuntimeError("ConvertPirAttribute2RuntimeAttribute not support")


class ExtractionDatabase:
    def __init__(self, owner_id, upload_root):
        self.owner_id = owner_id
        self.claim_id = uuid4()
        self.document_id = uuid4()
        self.document = {
            "document_id": self.document_id,
            "claim_id": self.claim_id,
            "document_type": "Repair Estimate",
            "document_url": "claims/test/estimate.txt",
            "original_file_name": "estimate.txt",
            "mime_type": "text/plain",
        }
        destination = upload_root / self.document["document_url"]
        destination.parent.mkdir(parents=True)
        destination.write_text("Repair estimate: EGP 7,500", encoding="utf-8")
        self.extraction = None
        self.committed = False

    def execute(self, statement, params=None):
        sql = str(statement)
        if "SELECT c.claim_id, p.user_id" in sql:
            return Result({"claim_id": self.claim_id, "user_id": self.owner_id} if str(params["claim_id"]) == str(self.claim_id) else None)
        if "FROM claim_documents WHERE document_id" in sql:
            return Result(self.document if str(params["document_id"]) == str(self.document_id) else None)
        if "FROM claim_extractions WHERE claim_id" in sql:
            return Result(self.extraction)
        if "INSERT INTO claim_extractions" in sql:
            self.extraction = {
                "extraction_id": uuid4(), "claim_id": self.claim_id,
                "extracted_data": json.loads(params["extracted_data"]),
                "extraction_confidence": params["extraction_confidence"],
                "extracted_at": datetime(2026, 8, 25, 12),
            }
            return Result(self.extraction)
        if "INSERT INTO audit_logs" in sql:
            return Result(None)
        raise AssertionError(f"Unexpected SQL: {sql}")

    def commit(self):
        self.committed = True

    def rollback(self):
        pass


class DocumentExtractionTests(unittest.TestCase):
    def test_structures_driver_licence_fields_and_normalizes_dates(self):
        data = extract_structured_data("Driver's Licence", """SOUTH CAROLINA
Name: SAMPLE CAROLINE
Licence Number: 1234567890
DOB: 08/27/2000
Address: 813 MAIN STREET, COLUMBIA, SC 29201
Class: C
Restrictions: NONE
Issue Date: 08/27/2025
Expiry Date: 08/27/2033""")
        self.assertEqual(data["full_name"], "SAMPLE CAROLINE")
        self.assertEqual(data["licence_number"], "1234567890")
        self.assertEqual(data["date_of_birth"], "2000-08-27")
        self.assertEqual(data["expiry_date"], "2033-08-27")
        self.assertEqual(data["issuing_authority"], "South Carolina")

    def test_structures_vehicle_registration_and_repair_estimate(self):
        registration = extract_structured_data("Vehicle Registration", """Owner: Mona Adel
Registration Number: ABC-1234
Plate Number: ABC-1234
Make: Toyota
Model: Corolla
Year: 2022
VIN: 1HGCM82633A004352
Registration Expiry: 2027-06-18""")
        self.assertEqual(registration["owner_name"], "Mona Adel")
        self.assertEqual(registration["model_year"], 2022)
        self.assertEqual(registration["registration_expiry_date"], "2027-06-18")

        estimate = extract_structured_data("Repair Estimate", """Garage: ABC Garage
Estimate Number: EST-102
Estimate Date: 2026-06-18
Customer: Mona Adel
Vehicle: Toyota Corolla
Damage: rear bumper deformation, trunk lid creasing
- Rear bumper replacement EGP 2850
- Paint work EGP 4650
Total Estimated: EGP 7,500""")
        self.assertEqual(estimate["garage_name"], "ABC Garage")
        self.assertEqual(estimate["issued_date"], "2026-06-18")
        self.assertEqual(estimate["total_estimated_repair_cost"], 7500.0)
        self.assertEqual(estimate["currency"], "EGP")
        self.assertEqual(len(estimate["repair_items"]), 0)

    def test_repair_estimate_label_value_layout_is_complete_and_never_uses_titles(self):
        data = extract_structured_data("Repair Estimate", """ABC Garage
Vehicle Repair Estimate

Estimate No.
EST-2026-0618

Issued
18 June 2026

Customer
Mona Adel

Vehicle
Toyota Corolla - 2022

Registration
ABC-1234

Claim reference
AXA-MTR-2026-00421

Damage assessment
Inspection found rear-end collision damage: rear bumper deformation, trunk lid creasing, and a cracked right tail lamp. The vehicle is safe to tow but repairs are required before normal use.

Estimated repair items

Item | Parts (EGP) | Labour (EGP) | Total (EGP)
Rear bumper cover replacement | 2,850 | 900 | 3,750
Trunk lid repair and paint | 1,250 | 1,100 | 2,350
Right tail lamp assembly | 980 | 120 | 1,100
Paint materials and alignment | 0 | 300 | 300

TOTAL ESTIMATED REPAIR COST
7,500

This estimate is valid for 30 days and is provided for insurance claim assessment.""")
        self.assertEqual(data["garage_name"], "ABC Garage")
        self.assertEqual(data["estimate_number"], "EST-2026-0618")
        self.assertEqual(data["issued_date"], "2026-06-18")
        self.assertEqual(data["customer_name"], "Mona Adel")
        self.assertEqual(data["vehicle"]["make_model"], "Toyota Corolla")
        self.assertEqual(data["vehicle"]["year"], 2022)
        self.assertEqual(data["registration_number"], "ABC-1234")
        self.assertEqual(data["claim_reference"], "AXA-MTR-2026-00421")
        self.assertEqual(data["total_estimated_repair_cost"], 7500.0)
        self.assertEqual(len(data["repair_items"]), 4)
        self.assertIn("rear-end collision damage", data["damage_description"])
        self.assertNotEqual(data["damage_description"], "assessment")
        self.assertNotEqual(data["vehicle"], "Repair Estimate")
        self.assertNotEqual(data["garage_name"], "Vehicle Repair Estimate")

    def test_visual_and_other_dedicated_parsers_do_not_invent_facts(self):
        visual = extract_structured_data("Photos of Damage", "")
        self.assertEqual(visual, {"visual_evidence_preserved": True, "structured_extraction_available": False})
        receipt = extract_structured_data("Receipts for Essentials", "Invoice Number\nREC-42\nTotal Amount\nEGP 250")
        self.assertEqual(receipt["invoice_number"], "REC-42")
        self.assertEqual(receipt["total_amount"], 250.0)
        self.assertIsNone(receipt["provider_name"])

    def test_vehicle_registration_uses_full_labels_and_validates_vin(self):
        registration = extract_structured_data("Vehicle Registration", """Registration Number
ABC-1234
Vehicle Identification Number (VIN)
JTDBR32E720123456
Owner Name
Mona Adel
Vehicle Make
Toyota
Vehicle Model
Corolla
Model Year
2022
Registration Expiry Date
31 May 2027""")
        self.assertEqual(registration["registration_number"], "ABC-1234")
        self.assertEqual(registration["vin"], "JTDBR32E720123456")
        self.assertEqual(registration["owner_name"], "Mona Adel")
        self.assertEqual(registration["vehicle_make"], "Toyota")
        self.assertEqual(registration["vehicle_model"], "Corolla")
        self.assertEqual(registration["registration_expiry_date"], "2027-05-31")

    def test_police_report_and_licence_validate_fields_and_preserve_narrative(self):
        police = extract_structured_data("Police Report", """REPORT NUMBER
PR-12345
INCIDENT DATE
18 June 2026
INCIDENT TIME
14:35
LOCATION
Nasr City, Cairo
DRIVER
Mona Adel
VEHICLE
Toyota Corolla - Registration ABC-1234
INCIDENT SUMMARY
Vehicle was hit from behind while stopped at a traffic signal.""")
        self.assertEqual(police["report_number"], "PR-12345")
        self.assertEqual(police["incident_date"], "2026-06-18")
        self.assertEqual(police["incident_time"], "14:35")
        self.assertEqual(police["incident_location"], "Nasr City, Cairo")
        self.assertEqual(police["driver_name"], "Mona Adel")
        self.assertEqual(police["registration_number"], "ABC-1234")
        self.assertIn("hit from behind", police["raw_narrative"])
        licence = extract_structured_data("Driver's Licence", "Class\n9aEnd\nLicence Number\nDL-10001")
        self.assertIsNone(licence["vehicle_class"])
        self.assertEqual(licence["licence_number"], "DL-10001")

    def test_ocr_style_labels_and_merged_values_use_the_same_document_parsers(self):
        police = extract_structured_data("Police Report", """POLICEREPORT
REPORTNUMBER
PR-12345
INCIDENTDATE
18 June 2026
INCIDENTTIME
14:35
LOCATION
Nasr City, Cairo
DRIVER
Mona Adel
VEHICLE
Toyota Corolla - Registration ABC-1234
INCIDENTSUMMARY
Vehicle was hit from behind while stopped at a traffic signal.
OFFICER NOTE
No injuries were reported at the scene.""", "image_ocr")
        self.assertEqual(police["report_number"], "PR-12345")
        self.assertEqual(police["incident_date"], "2026-06-18")
        self.assertEqual(police["incident_time"], "14:35")
        self.assertEqual(police["registration_number"], "ABC-1234")
        self.assertFalse(police["injuries_reported"])
        self.assertIn("No injuries", police["officer_notes"])

        licence = extract_structured_data("Driver's Licence", """A1234567890
3Date of bith
01/01/2000
F
8123MAINSTREET
COLUMBIA,SC29201
Class
9aEnd
D
NONE
4alssue
4bExp
08/27/202508/27/2033""", "scanned_pdf_ocr")
        self.assertEqual(licence["licence_number"], "A1234567890")
        self.assertEqual(licence["date_of_birth"], "2000-01-01")
        self.assertEqual(licence["sex"], "F")
        self.assertEqual(licence["vehicle_class"], "D")
        self.assertEqual(licence["issue_date"], "2025-08-27")
        self.assertEqual(licence["expiry_date"], "2033-08-27")

        registration = extract_structured_data("Vehicle Registration", "Engine Capacity\n1800 CC")
        self.assertEqual(registration["engine_capacity_cc"], 1800)

    def test_rejects_invalid_typed_ocr_values_and_extracts_valid_sex(self):
        licence = extract_structured_data("Driver's Licence", """15Sex
01/01/2000
F
Date of bith
01/01/2000
Class
9aEnd""", "image_ocr")
        self.assertEqual(licence["sex"], "F")
        self.assertEqual(licence["date_of_birth"], "2000-01-01")
        self.assertIsNone(licence["vehicle_class"])

        registration = extract_structured_data("Vehicle Registration", """Engine Capacity
not 1800 cc
Fuel Type
petro1
Number of Seats
five""", "image_ocr")
        self.assertIsNone(registration["engine_capacity_cc"])
        self.assertIsNone(registration["fuel_type"])
        self.assertIsNone(registration["number_of_seats"])

    def test_reconstructs_wrapped_text_until_the_next_recognised_label(self):
        police = extract_structured_data("Police Report", """INCIDENTSUMMARY
Vehicle was hit from behind while stopped at a traffic signal. Visible damage includes
the rear bumper, trunk lid, and right tail lamp.
OFFICER NOTE
No injuries were reported at the scene. Parties were advised to submit this report to
their insurer.
REPORT NUMBER
PR-12345""", "image_ocr")
        self.assertEqual(
            police["incident_summary"],
            "Vehicle was hit from behind while stopped at a traffic signal. Visible damage includes the rear bumper, trunk lid, and right tail lamp.",
        )
        self.assertEqual(
            police["officer_notes"],
            "No injuries were reported at the scene. Parties were advised to submit this report to their insurer.",
        )
        self.assertEqual(police["report_number"], "PR-12345")
        self.assertFalse(police["injuries_reported"])

    def test_shared_text_parser_reconstructs_health_and_travel_descriptions(self):
        medical = extract_structured_data("Medical Report", """Diagnosis
Sprained ankle after a fall while travelling.
Treatment
Ankle immobilisation and pain relief were prescribed over
the following seven days.
Report Date
18 June 2026""")
        self.assertEqual(medical["diagnosis"], "Sprained ankle after a fall while travelling.")
        self.assertEqual(medical["treatment"], "Ankle immobilisation and pain relief were prescribed over the following seven days.")
        self.assertEqual(medical["report_date"], "2026-06-18")

    def test_every_requirement_document_type_has_an_intentional_parser_or_visual_result(self):
        from app.claim_requirements import REQUIRED_DOCUMENTS

        document_types = {document_type for claim_types in REQUIRED_DOCUMENTS.values() for documents in claim_types.values() for document_type in documents}
        self.assertEqual(len(document_types), 28)
        for document_type in document_types:
            with self.subTest(document_type=document_type):
                text = "Garage: Example Garage\nEstimate Number: EST-100" if document_type == "Repair Estimate" else "Reference: REF-100\nDate: 18 June 2026"
                result = extract_structured_data(document_type, text)
                if document_type in {"Photos of Damage", "Spare Key"}:
                    self.assertEqual(result["visual_evidence_preserved"], True)
                    self.assertEqual(result["structured_extraction_available"], False)
                else:
                    self.assertNotEqual(result, {})

    def test_structures_invoice_medical_police_and_member_id_without_fabrication(self):
        invoice = extract_structured_data("Itemised Invoice", """Provider: Cairo Clinic
Invoice Number: INV-77
Invoice Date: 2026-06-18
Subtotal: EGP 1000
VAT: EGP 140
Total Amount: EGP 1140""")
        self.assertEqual(invoice["total_amount"], 1140.0)
        self.assertEqual(invoice["currency"], "EGP")
        medical = extract_structured_data("Medical Report", "Patient: Layla Mostafa\nDiagnosis: Sprained ankle\nReport Date: 2026-06-18")
        self.assertEqual(medical["diagnosis"], "Sprained ankle")
        police = extract_structured_data("Police Report", "Report Number: PR-100\nIncident Date: 2026-06-18\nIncident Location: Nasr City")
        self.assertEqual(police["report_number"], "PR-100")
        member = extract_structured_data("Member ID", "Member Name: Layla Mostafa\nMember ID: MID-123\nPolicy Number: POL-88\nExpiry: 2027-06-18")
        self.assertEqual(member["member_id"], "MID-123")
        self.assertEqual(extract_structured_data("Repair Estimate", "unreadable OCR fragments")["structured_extraction_available"], False)
    def test_extracts_native_docx_text(self):
        with TemporaryDirectory() as directory:
            path = Path(directory) / "estimate.docx"
            document = Document()
            document.add_paragraph("Garage repair estimate: EGP 7,500")
            document.save(path)
            result = extract_document_content(path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "Repair Estimate")
        self.assertEqual(result.strategy, "native_docx_text")
        self.assertIn("EGP 7,500", result.text)
        self.assertEqual(result.confidence, 1.0)

    def test_extracts_native_text_without_ocr(self):
        with TemporaryDirectory() as directory:
            path = Path(directory) / "estimate.txt"
            path.write_text("Repair estimate: EGP 7,500", encoding="utf-8")
            result = extract_document_content(path, "text/plain", "Repair Estimate")
        self.assertEqual(result.strategy, "native_text")
        self.assertEqual(result.text, "Repair estimate: EGP 7,500")
        self.assertEqual(result.confidence, 1.0)

    def test_native_pdf_text_bypasses_document_ocr(self):
        with patch("app.document_extraction._extract_pdf_text", return_value="Report Number: PR-100") as native, patch("app.document_extraction._extract_scanned_pdf_text") as scanned:
            result = extract_document_content(Path("report.pdf"), "application/pdf", "Police Report")
        self.assertEqual(result.strategy, "native_pdf_text")
        self.assertEqual(result.text, "Report Number: PR-100")
        native.assert_called_once()
        scanned.assert_not_called()

    def _legacy_azure_image_ocr_structure_test(self):
        page = SimpleNamespace(
            page_number=1,
            angle=None,
            lines=[
                SimpleNamespace(content="رقم الوثيقة", polygon=[SimpleNamespace(x=10, y=5), SimpleNamespace(x=90, y=5), SimpleNamespace(x=90, y=20), SimpleNamespace(x=10, y=20)]),
                SimpleNamespace(content="Policy Number: POL-88", polygon=[SimpleNamespace(x=10, y=30), SimpleNamespace(x=180, y=30), SimpleNamespace(x=180, y=45), SimpleNamespace(x=10, y=45)]),
            ],
        )
        table = SimpleNamespace(
            row_count=2,
            column_count=2,
            bounding_regions=[SimpleNamespace(page_number=1, polygon=[SimpleNamespace(x=10, y=50), SimpleNamespace(x=150, y=50), SimpleNamespace(x=150, y=90), SimpleNamespace(x=10, y=90)])],
            cells=[
                SimpleNamespace(row_index=0, column_index=0, row_span=1, column_span=1, kind="columnHeader", content="Item"),
                SimpleNamespace(row_index=0, column_index=1, row_span=1, column_span=1, kind="columnHeader", content="Cost"),
                SimpleNamespace(row_index=1, column_index=0, row_span=1, column_span=1, kind="content", content="Bumper"),
                SimpleNamespace(row_index=1, column_index=1, row_span=1, column_span=1, kind="content", content="EGP 100"),
            ],
        )
        paragraph = SimpleNamespace(role="title", bounding_regions=[SimpleNamespace(page_number=1, polygon=[SimpleNamespace(x=10, y=5), SimpleNamespace(x=90, y=5), SimpleNamespace(x=90, y=20), SimpleNamespace(x=10, y=20)])])
        azure = StubAzureClient(SimpleNamespace(pages=[page], paragraphs=[paragraph], tables=[table]))
        with TemporaryDirectory() as directory:
            path = Path(directory) / "mixed.jpg"
            Image.new("RGB", (20, 20), "white").save(path)
            with patch("app.document_extraction._azure_document_intelligence_client", return_value=azure), patch("app.document_extraction._ocr_pipelines") as paddle:
                result = extract_document_content(path, "image/jpeg", "Member ID")
        self.assertEqual(azure.calls[0][0], "prebuilt-layout")
        paddle.assert_not_called()
        self.assertEqual(result.strategy, "image_ocr")
        self.assertEqual(result.text.splitlines(), ["رقم الوثيقة", "Policy Number: POL-88"])
        self.assertIsNone(result.confidence)
        self.assertEqual(result.structure["pages"][0]["blocks"][0]["bbox"], [10.0, 5.0, 90.0, 20.0])
        self.assertEqual(result.structure["pages"][0]["blocks"][0]["language_pass"], "azure")
        self.assertEqual(result.structure["pages"][0]["layout"][0]["type"], "title")
        self.assertIn("<th>Item</th>", result.structure["pages"][0]["tables"][0]["html"])
        self.assertEqual(result.structure["pages"][0]["tables"][0]["bbox"], [10.0, 50.0, 150.0, 90.0])

    def _legacy_azure_image_ocr_text_test(self):
        page = SimpleNamespace(page_number=1, angle=0, lines=[SimpleNamespace(content="English text", confidence=0.98, polygon=[SimpleNamespace(x=0, y=0), SimpleNamespace(x=100, y=0), SimpleNamespace(x=100, y=10), SimpleNamespace(x=0, y=10)])])
        azure = StubAzureClient(SimpleNamespace(pages=[page], paragraphs=[], tables=[]))
        with TemporaryDirectory() as directory:
            path = Path(directory) / "document.png"
            Image.new("RGB", (20, 20), "white").save(path)
            with patch("app.document_extraction._azure_document_intelligence_client", return_value=azure):
                result = extract_document_content(path, "image/png", "Member ID")
        self.assertEqual(result.text, "English text")
        self.assertEqual(result.confidence, 0.98)

    def _legacy_azure_table_html_test(self):
        from app.document_extraction import _azure_table_html

        table = SimpleNamespace(
            row_count=2,
            column_count=2,
            cells=[SimpleNamespace(row_index=0, column_index=0, row_span=2, column_span=1, kind="content", content="Merged"), SimpleNamespace(row_index=0, column_index=1, row_span=1, column_span=1, kind="content", content="Top"), SimpleNamespace(row_index=1, column_index=1, row_span=1, column_span=1, kind="content", content="Bottom")],
        )
        self.assertIn('<td rowspan="2">Merged</td>', _azure_table_html(table))

    def _legacy_azure_missing_configuration_test(self):
        from app.document_extraction import _azure_document_intelligence_client, settings

        _azure_document_intelligence_client.cache_clear()
        with TemporaryDirectory() as directory, patch.object(settings, "AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT", None), patch.object(settings, "AZURE_DOCUMENT_INTELLIGENCE_KEY", None):
            path = Path(directory) / "document.png"
            Image.new("RGB", (20, 20), "white").save(path)
            with self.assertRaisesRegex(DocumentExtractionError, "not configured"):
                extract_document_content(path, "image/png", "Member ID")
        _azure_document_intelligence_client.cache_clear()

    def _legacy_azure_failure_test(self):
        with TemporaryDirectory() as directory:
            path = Path(directory) / "document.png"
            Image.new("RGB", (20, 20), "white").save(path)
            with patch("app.document_extraction._azure_document_intelligence_client", side_effect=RuntimeError("service unavailable")):
                with self.assertLogs("app.document_extraction", level="ERROR") as logs:
                    with self.assertRaisesRegex(DocumentExtractionError, "OCR could not read the image"):
                        extract_document_content(path, "image/png", "Member ID")
        self.assertIn("service unavailable", "\n".join(logs.output))

    def test_ocr_disables_onednn_before_pipeline_initialization(self):
        with patch.dict(os.environ, {}, clear=True):
            from app.document_extraction import _disable_onednn

            _disable_onednn()
            self.assertEqual(os.environ["FLAGS_use_mkldnn"], "0")

    def test_paddle_document_pipeline_processes_scanned_pdf(self):
        sample = Path(__file__).resolve().parents[1] / "data" / "AXA_capstone_data" / "sample_documents" / "CLM-001.jpg"
        english = StubStructurePipeline({"res": {"overall_ocr_res": {"rec_texts": ["CLM-001"], "rec_scores": [0.9], "rec_polys": [[[0, 0], [100, 0], [100, 20], [0, 20]]]}}})
        arabic = StubStructurePipeline({"res": {"overall_ocr_res": {"rec_texts": [], "rec_scores": [], "rec_polys": []}}})
        with TemporaryDirectory() as directory:
            path = Path(directory) / "scan.pdf"
            with Image.open(sample) as image:
                image.convert("RGB").save(path, "PDF")
            with patch("app.document_extraction._ocr_pipelines", return_value=(("english", english), ("arabic", arabic))):
                result = extract_document_content(path, "application/pdf", "Repair Estimate")
        self.assertEqual(result.strategy, "scanned_pdf_ocr")
        self.assertIn("CLM-001", result.text)

    @unittest.skip("Image OCR uses external vision providers; this test covers only the scanned-PDF pipeline")
    def test_paddle_document_pipeline_preserves_arabic_and_mixed_text_reading_order(self):
        english = StubStructurePipeline({"res": {"overall_ocr_res": {"rec_texts": ["Policy Number: POL-88"], "rec_scores": [0.95], "rec_polys": [[[0, 30], [200, 30], [200, 50], [0, 50]]]}}})
        arabic = StubStructurePipeline({"res": {"overall_ocr_res": {"rec_texts": ["رقم الوثيقة"], "rec_scores": [0.93], "rec_polys": [[[0, 0], [150, 0], [150, 20], [0, 20]]]}}})
        with TemporaryDirectory() as directory:
            path = Path(directory) / "mixed.png"
            Image.new("RGB", (20, 20), "white").save(path)
            with patch("app.document_extraction._ocr_pipelines", return_value=(("english", english), ("arabic", arabic))):
                result = extract_document_content(path, "image/png", "Member ID")
        self.assertEqual(result.text.splitlines(), ["رقم الوثيقة", "Policy Number: POL-88"])
        self.assertEqual(result.structure["reading_order"], "top_to_bottom_left_to_right")

    def test_visual_evidence_provider_failure_is_retained_for_safe_validation(self):
        result = extract_document_content(Path("photo.jpg"), "image/jpeg", "Photos of Damage")
        self.assertEqual(result.strategy, "visual_content_analysis")
        self.assertEqual(result.text, "")
        self.assertIsNone(result.confidence)

    def test_rejects_unsupported_files(self):
        with self.assertRaises(DocumentExtractionError):
            extract_document_content(Path("estimate.xlsx"), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "Repair Estimate")

    def test_authorized_processing_links_extraction_to_document_and_is_idempotent(self):
        owner_id = uuid4()
        with TemporaryDirectory() as directory, patch("app.main.settings.UPLOAD_DIR", directory):
            db = ExtractionDatabase(owner_id, Path(directory))
            first = extract_claim_document(str(db.claim_id), str(db.document_id), {"user_id": owner_id, "role_name": "Customer"}, db)
            second = extract_claim_document(str(db.claim_id), str(db.document_id), {"user_id": owner_id, "role_name": "Customer"}, db)
        self.assertFalse(first["reused"])
        self.assertTrue(second["reused"])
        self.assertEqual(first["document_id"], str(db.document_id))
        self.assertEqual(db.extraction["extracted_data"]["document_id"], str(db.document_id))
        self.assertEqual(db.extraction["extracted_data"]["extracted_text"], "Repair estimate: EGP 7,500")
        self.assertEqual(db.extraction["extracted_data"]["raw_extraction"]["text"], "Repair estimate: EGP 7,500")
        self.assertEqual(db.extraction["extracted_data"]["structured_data"]["structured_extraction_available"], False)
        self.assertEqual(first["raw_text"], "Repair estimate: EGP 7,500")
        self.assertEqual(first["validation"]["status"], "valid")
        self.assertTrue(first["validation"]["document_valid"])
        self.assertEqual(first["structured_data"]["structured_extraction_available"], False)
        self.assertTrue(db.committed)

    def test_extraction_response_reports_a_wrong_document_type(self):
        owner_id = uuid4()
        with TemporaryDirectory() as directory, patch("app.main.settings.UPLOAD_DIR", directory):
            db = ExtractionDatabase(owner_id, Path(directory))
            db.document["document_type"] = "Repair Estimate"
            destination = Path(directory) / db.document["document_url"]
            destination.write_text("DRIVER'S LICENCE\nLicence Number: DL-10001\nDate of Birth: 01/01/2000", encoding="utf-8")
            response = extract_claim_document(str(db.claim_id), str(db.document_id), {"user_id": owner_id, "role_name": "Customer"}, db)
        self.assertEqual(response["validation"]["status"], "invalid")
        self.assertFalse(response["validation"]["document_valid"])
        self.assertEqual(response["validation"]["expected_document_type"], "Repair Estimate")
        self.assertEqual(response["validation"]["detected_document_type"], "Driver's Licence")

    def test_rejects_another_customers_claim_and_missing_files(self):
        owner_id, other_id = uuid4(), uuid4()
        with TemporaryDirectory() as directory, patch("app.main.settings.UPLOAD_DIR", directory):
            db = ExtractionDatabase(owner_id, Path(directory))
            with self.assertRaises(HTTPException) as unauthorized:
                extract_claim_document(str(db.claim_id), str(db.document_id), {"user_id": other_id, "role_name": "Customer"}, db)
            self.assertEqual(unauthorized.exception.status_code, 403)
            (Path(directory) / db.document["document_url"]).unlink()
            with self.assertRaises(HTTPException) as missing:
                extract_claim_document(str(db.claim_id), str(db.document_id), {"user_id": owner_id, "role_name": "Customer"}, db)
            self.assertEqual(missing.exception.status_code, 409)
            self.assertIsNone(db.extraction)

    def test_extraction_failure_does_not_create_a_record(self):
        owner_id = uuid4()
        with TemporaryDirectory() as directory, patch("app.main.settings.UPLOAD_DIR", directory):
            db = ExtractionDatabase(owner_id, Path(directory))
            with patch("app.main.extract_document_content", side_effect=DocumentExtractionError("OCR could not read the document")):
                with self.assertRaises(HTTPException) as failed:
                    extract_claim_document(str(db.claim_id), str(db.document_id), {"user_id": owner_id, "role_name": "Customer"}, db)
        self.assertEqual(failed.exception.status_code, 422)
        self.assertIsNone(db.extraction)

    def test_image_ocr_uses_openrouter_before_groq(self):
        from app.document_extraction import ExtractedDocument

        with TemporaryDirectory() as directory:
            path = Path(directory) / "document.jpg"
            Image.new("RGB", (20, 20), "white").save(path)
            expected = ExtractedDocument("image_ocr", "Policy Number: POL-88", None)
            with patch("app.document_extraction._openrouter_image_ocr", return_value=expected) as openrouter, patch("app.document_extraction._groq_image_ocr") as groq:
                result = extract_document_content(path, "image/jpeg", "Member ID")
        self.assertEqual(result, expected)
        openrouter.assert_called_once_with(path, "image_ocr")
        groq.assert_not_called()

    def test_openrouter_failure_and_timeout_use_groq_fallback(self):
        from app.document_extraction import ExtractedDocument

        for failure in (RuntimeError("provider error"), TimeoutError("timeout")):
            with self.subTest(failure=failure), TemporaryDirectory() as directory:
                path = Path(directory) / "document.png"
                Image.new("RGB", (20, 20), "white").save(path)
                with patch("app.document_extraction._openrouter_image_ocr", side_effect=failure), patch("app.document_extraction._groq_image_ocr", return_value=ExtractedDocument("image_ocr", "Fallback text", None)) as groq:
                    result = extract_document_content(path, "image/png", "Member ID")
                self.assertEqual(result.text, "Fallback text")
                groq.assert_called_once_with(path, "image_ocr")

    def test_empty_openrouter_result_uses_groq_fallback(self):
        from app.document_extraction import ExtractedDocument

        with TemporaryDirectory() as directory:
            path = Path(directory) / "document.webp"
            Image.new("RGB", (20, 20), "white").save(path, "WEBP")
            with patch("app.document_extraction._openrouter_image_ocr", side_effect=DocumentExtractionError("empty OCR result")), patch("app.document_extraction._groq_image_ocr", return_value=ExtractedDocument("image_ocr", "Fallback text", None)):
                result = extract_document_content(path, "image/webp", "Member ID")
        self.assertEqual(result.text, "Fallback text")

    def test_image_ocr_raises_when_both_providers_fail(self):
        with TemporaryDirectory() as directory:
            path = Path(directory) / "document.png"
            Image.new("RGB", (20, 20), "white").save(path)
            with patch("app.document_extraction._openrouter_image_ocr", side_effect=RuntimeError("primary failed")), patch("app.document_extraction._groq_image_ocr", side_effect=RuntimeError("fallback failed")):
                with self.assertRaisesRegex(DocumentExtractionError, "OCR could not read the image"):
                    extract_document_content(path, "image/png", "Member ID")

    def test_vision_response_preserves_the_image_ocr_response_shape(self):
        from app.document_extraction import _vision_response_to_extracted_document

        result = _vision_response_to_extracted_document(json.dumps({
            "text": "رقم الوثيقة\nPolicy Number: POL-88",
            "blocks": [{"type": "text", "text": "رقم الوثيقة", "language": "arabic", "bbox": None}],
            "layout": [{"type": "title", "text": "رقم الوثيقة"}],
            "tables": [{"html": "<table><tr><td>POL-88</td></tr></table>", "rows": [["POL-88"]]}],
        }), "image_ocr", "openrouter")
        self.assertEqual(result.strategy, "image_ocr")
        self.assertEqual(result.text, "رقم الوثيقة\nPolicy Number: POL-88")
        self.assertIsNone(result.structure["pages"][0]["blocks"][0]["bbox"])
        self.assertEqual(result.structure["pages"][0]["layout"][0]["type"], "title")

    def test_vision_response_carries_model_structured_data(self):
        from app.document_extraction import _vision_response_to_extracted_document

        result = _vision_response_to_extracted_document(json.dumps({
            "text": "Patient: Mona Adel",
            "structured_data": {"patient_name": "Mona Adel", "report_reference": "MED-2026-0610-014"},
        }), "image_ocr", "openrouter")
        self.assertEqual(result.structured_data["patient_name"], "Mona Adel")

    def test_docx_text_is_sent_to_ai_structured_extraction(self):
        from app.document_extraction import extract_document_content

        with TemporaryDirectory() as directory:
            path = Path(directory) / "medical.docx"
            document = Document()
            document.add_paragraph("MEDICAL REPORT\nPatient | Mona Adel\nReport reference | MED-2026-0610-014")
            document.save(path)
            expected = {"patient_name": "Mona Adel", "report_reference": "MED-2026-0610-014", "diagnosis": None}
            with patch("app.document_extraction._extract_structured_data_with_ai", return_value=expected) as structured:
                result = extract_document_content(path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "Repair Estimate")
        self.assertEqual(result.strategy, "native_docx_text")
        self.assertEqual(result.structured_data, expected)
        structured.assert_called_once()

    def test_ai_structured_data_reaches_the_existing_response(self):
        from app.document_extraction import ExtractedDocument

        owner_id = uuid4()
        with TemporaryDirectory() as directory, patch("app.main.settings.UPLOAD_DIR", directory):
            db = ExtractionDatabase(owner_id, Path(directory))
            extracted = ExtractedDocument("native_text", "MEDICAL REPORT", 1.0, structured_data={"patient_name": "Mona Adel"})
            with patch("app.main.extract_document_content", return_value=extracted):
                response = extract_claim_document(str(db.claim_id), str(db.document_id), {"user_id": owner_id, "role_name": "Customer"}, db)
        self.assertEqual(response["structured_data"], {"patient_name": "Mona Adel"})

    def test_structured_extraction_falls_back_to_groq_and_keeps_text_on_failure(self):
        from app.document_extraction import _extract_structured_data_with_ai

        with patch("app.document_extraction._openrouter_structured_extraction", side_effect=RuntimeError("primary failed")), patch("app.document_extraction._groq_structured_extraction", return_value={"patient_name": "Mona Adel"}) as groq:
            result = _extract_structured_data_with_ai("Patient: Mona Adel", "Medical Report")
        self.assertEqual(result, {"patient_name": "Mona Adel"})
        groq.assert_called_once()

        with patch("app.document_extraction._openrouter_structured_extraction", side_effect=RuntimeError("primary failed")), patch("app.document_extraction._groq_structured_extraction", side_effect=RuntimeError("fallback failed")):
            self.assertIsNone(_extract_structured_data_with_ai("Patient: Mona Adel", "Medical Report"))


if __name__ == "__main__":
    unittest.main()
